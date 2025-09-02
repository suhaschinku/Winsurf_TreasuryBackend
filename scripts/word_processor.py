import os
import time
from dotenv import load_dotenv
from langchain_community.document_loaders import Docx2txtLoader, UnstructuredWordDocumentLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from gen_ai_hub.proxy.langchain.init_models import init_embedding_model
from langchain.docstore.document import Document
import docx
from docx.document import Document as DocxDocument
from docx.table import Table
import mammoth
from logger_setup import get_logger
import re
from concurrent.futures import ThreadPoolExecutor
from env_config import TABLE_NAMES, EMBEDDING_MODEL, HANA_DB_API
from destination_srv import get_destination_service_credentials, generate_token, fetch_destination_details,extract_aicore_credentials
import requests
logger = get_logger()
load_dotenv()


#$$$ SOC: 28.05.25 -- Initialize AIC Credentials --- $$$#
logger.info ("====> pdf_processor.py -> AIC CREDENTIALS <====")

# Load VCAP_SERVICES from environment
vcap_services = os.environ.get("VCAP_SERVICES")

# Extract destination service credentials
destination_service_credentials = get_destination_service_credentials(vcap_services)

# Generate OAuth token for destination service
try:
    oauth_token = generate_token(
        uri=destination_service_credentials['dest_auth_url'] + "/oauth/token",
        client_id=destination_service_credentials['clientid'],
        client_secret=destination_service_credentials['clientsecret']
    )
except requests.exceptions.HTTPError as e:
    # Handle HTTP 500 error for invalid client secret
    if e.response is not None and e.response.status_code == 500:
        raise Exception("HTTP 500: Check if the client secret is correct.") from e
    else:
        raise


#-------------------------------- READ AIC Configuration -------------------------------------

# variables for AIC credentials
global AIC_CREDENTIALS

# Get AIC details from Dest Services
dest_AIC = "GENAI_AI_CORE"
aicore_details = fetch_destination_details(
    destination_service_credentials['dest_base_url'],
    dest_AIC,
    oauth_token
)
    
# Extract AIC Details
AIC_CREDENTIALS = extract_aicore_credentials(aicore_details)
logger.info("pdf_processor -> AIC Credential", AIC_CREDENTIALS)

#$$$ EOC: 28.05.25 -- Initialize AIC Credentials --- $$$#




def clean_text(text):
    """Clean OCR artifacts and normalize text."""
    if not text:
        return ""
    text = re.sub(r"JPHORGAN", "JPMORGAN", text, flags=re.IGNORECASE)
    text = re.sub(r"excape", "except", text, flags=re.IGNORECASE)
    text = re.sub(r"\s+", " ", text).strip()  # Normalize whitespace
    return text

def table_to_text(table, source_file, table_idx, source_tool="python-docx"):
    """Convert Word table data to natural language text with improved header detection."""
    if not table:
        return ""
    
    try:
        # Extract table data from docx table
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data.append(row_data)
        
        if len(table_data) < 2:  # Need at least header + 1 row
            return ""
        
        text = []
        headers = table_data[0]
        
        # Check if first row contains only numbers (not a header)
        if all(re.match(r"^\$?\d[\d,.]*$", str(cell).strip()) for cell in headers if str(cell).strip()):
            headers = ["Column " + str(i) for i in range(len(headers))]  # Fallback if no clear header
        else:
            table_data = table_data[1:]  # Skip header row
        
        for row in table_data:
            if len(row) != len(headers):
                continue  # Skip malformed rows
            row_text = " ".join(f"{headers[i]} is {cell}" for i, cell in enumerate(row) if cell and str(cell).strip())
            if row_text:
                text.append(row_text)
        
        return "\n".join(text) if text else ""
    except Exception as e:
        logger.warning(f"Error processing table in {source_file}: {e}")
        return ""

def process_docx_with_python_docx(doc_path):
    """Extract text and tables using python-docx library."""
    try:
        doc = docx.Document(doc_path)
        filename = os.path.basename(doc_path)
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        # Combine paragraphs into text
        text_content = "\n".join(paragraphs)
        
        # Create text document
        text_docs = []
        if text_content:
            metadata = {
                "source_file": filename,
                "type": "text",
                "extractor": "python-docx"
            }
            text_docs.append(Document(page_content=clean_text(text_content), metadata=metadata))
        
        # Extract tables
        table_docs = []
        for table_idx, table in enumerate(doc.tables):
            table_text = table_to_text(table, filename, table_idx, source_tool="python-docx")
            if table_text:
                metadata = {
                    "source_file": filename,
                    "table_idx": table_idx,
                    "type": "table",
                    "extractor": "python-docx"
                }
                table_docs.append(Document(page_content=clean_text(table_text), metadata=metadata))
        
        return text_docs + table_docs
    except Exception as e:
        logger.warning(f"python-docx failed for {doc_path}: {e}")
        return []

def process_docx_with_mammoth(doc_path):
    """Extract text using mammoth library (good for .doc files)."""
    try:
        filename = os.path.basename(doc_path)
        
        with open(doc_path, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            text_content = result.value
        
        if text_content:
            metadata = {
                "source_file": filename,
                "type": "text",
                "extractor": "mammoth"
            }
            return [Document(page_content=clean_text(text_content), metadata=metadata)]
        return []
    except Exception as e:
        logger.warning(f"mammoth failed for {doc_path}: {e}")
        return []

def process_docx_with_langchain(doc_path):
    """Extract text using LangChain loaders."""
    try:
        filename = os.path.basename(doc_path)
        
        # Try Docx2txtLoader first
        try:
            loader = Docx2txtLoader(doc_path)
            documents = loader.load()
            if documents:
                for doc in documents:
                    doc.metadata['source_file'] = filename
                    doc.metadata['extractor'] = 'docx2txt'
                    doc.page_content = clean_text(doc.page_content)
                return documents
        except Exception as e:
            logger.warning(f"Docx2txtLoader failed for {doc_path}: {e}")
        
        # Fallback to UnstructuredWordDocumentLoader
        try:
            loader = UnstructuredWordDocumentLoader(doc_path)
            documents = loader.load()
            if documents:
                for doc in documents:
                    doc.metadata['source_file'] = filename
                    doc.metadata['extractor'] = 'unstructured'
                    doc.page_content = clean_text(doc.page_content)
                return documents
        except Exception as e:
            logger.warning(f"UnstructuredWordDocumentLoader failed for {doc_path}: {e}")
        
        return []
    except Exception as e:
        logger.warning(f"LangChain loaders failed for {doc_path}: {e}")
        return []

def process_doc(doc_path):
    """Extract text and tables from a DOC/DOCX file, returning chunks."""
    logger.info(f"Processing DOC/DOCX: {doc_path}")
    try:
        filename = os.path.basename(doc_path)
        all_docs = []
        
        # Method 1: Use python-docx for .docx files (best for tables)
        if doc_path.lower().endswith('.docx'):
            docs = process_docx_with_python_docx(doc_path)
            if docs:
                all_docs.extend(docs)
                logger.info(f"python-docx extracted {len(docs)} elements from {filename}")
        
        # Method 2: Use mammoth as fallback (good for .doc files)
        if not all_docs:
            docs = process_docx_with_mammoth(doc_path)
            if docs:
                all_docs.extend(docs)
                logger.info(f"mammoth extracted {len(docs)} elements from {filename}")
        
        # Method 3: Use LangChain loaders as final fallback
        if not all_docs:
            docs = process_docx_with_langchain(doc_path)
            if docs:
                all_docs.extend(docs)
                logger.info(f"LangChain extracted {len(docs)} elements from {filename}")
        
        if not all_docs:
            logger.warning(f"No content extracted from {filename}")
            return []
        
        # Step 2: Split into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        chunks = text_splitter.split_documents(all_docs)
        
        if not chunks:
            logger.warning(f"No chunks created from {filename}")
            return []
        
        # Append filename to each chunk's content
        for chunk in chunks:
            chunk.page_content = f"{chunk.page_content}\n\nSource: {filename}"
        
        logger.info(f"Processed {len(chunks)} chunks from {filename}")
        return chunks
    except Exception as e:
        logger.error(f"Error processing {doc_path}: {str(e)}", exc_info=True)
        return []

def create_embeddings(docs, model_name=EMBEDDING_MODEL, batch_size=400):
    """Generate embeddings for document chunks in batches."""
    if not docs:
        logger.warning("No documents provided for embedding")
        return [], 1
    
    logger.info(f"Creating embeddings for {len(docs)} docs with {model_name}")
    try:
        from gen_ai_hub.proxy import GenAIHubProxyClient
        logger.info("PDF Processor: AIC, {AIC_CREDENTIALS}")

        proxy_client = GenAIHubProxyClient(
                                    base_url = AIC_CREDENTIALS['aic_base_url'],
                                    auth_url = AIC_CREDENTIALS['aic_auth_url'],
                                    client_id = AIC_CREDENTIALS['clientid'],
                                    client_secret = AIC_CREDENTIALS['clientsecret'],
                                    resource_group = AIC_CREDENTIALS['resource_group']
                                    )

        embedding_model = init_embedding_model(model_name = EMBEDDING_MODEL, proxy_client=proxy_client)
    ### EOC: Initialize Embedding Models ### 
       
        logger.info(f"Embedding model initialized: {embedding_model}")
        
        results = []
        for i in range(0, len(docs), batch_size):
            batch = docs[i:i + batch_size]
            batch_texts = [doc.page_content for doc in batch]
            
            if not batch_texts or any(not isinstance(text, str) for text in batch_texts):
                logger.error(f"Invalid batch content at index {i}: {batch_texts[:50]}")
                return [], 1
            
            logger.info(f"Processing batch {i//batch_size + 1}, size: {len(batch_texts)}, Sample: {batch_texts[0][:50]}")
            embeddings = embedding_model.embed_documents(batch_texts)
            
            if not embeddings or len(embeddings) != len(batch_texts):
                logger.error(f"Embedding mismatch: {len(embeddings)} embeddings for {len(batch_texts)} texts")
                return [], 1
            
            results.extend(zip(batch, embeddings))
            logger.info(f"Batch {i//batch_size + 1} completed")
            
            if i + batch_size < len(docs):
                time.sleep(10)  # Rate limiting delay
        
        logger.info(f"Created embeddings for {len(results)} docs")
        return results, 0
    except Exception as e:
        logger.error(f"Embedding error: {str(e)}", exc_info=True)
        return [], 1

def process_doc_with_embeddings(doc_path, model=EMBEDDING_MODEL):
    """Process a single DOC/DOCX and generate embeddings."""
    chunks = process_doc(doc_path)
    if not chunks:
        logger.warning(f"No chunks to embed for {doc_path}")
        return []
    
    filename = os.path.basename(doc_path)
    logger.info(f"Using model {model} for {filename}")
    
    embeddings, error = create_embeddings(chunks, model)
    if error:
        print(f"❌ Embedding creation failed for {filename}!")
    else:
        print(f"✅ Created {len(embeddings)} embeddings for {filename}")
    
    return embeddings

def process_all_docs(directory, model=EMBEDDING_MODEL):
    """Process all DOC/DOCX files in a directory, separating transcript and non-transcript embeddings."""
    if not os.path.exists(directory):
        logger.error(f"Directory not found: {directory}")
        print(f"Directory not found: {directory}")
        return [], []
    
    doc_files = [f for f in os.listdir(directory) if f.lower().endswith(('.doc', '.docx'))]
    if not doc_files:
        print(f"No DOC/DOCX files in {directory}")
        return [], []
    
    print(f"Found {len(doc_files)} DOC/DOCX files")
    
    transcript_embeddings = []
    non_transcript_embeddings = []

    def process_single_doc(doc):
        """Process a single DOC/DOCX file."""
        doc_path = os.path.join(directory, doc)
        logger.info(f"Processing DOC/DOCX: {doc}")
        filename = os.path.basename(doc_path)
        is_transcript = 'transcript' in filename.lower()
        embeddings = process_doc_with_embeddings(doc_path, model)
        return embeddings, is_transcript

    # Use ThreadPoolExecutor for parallel processing
    with ThreadPoolExecutor(max_workers=4) as executor:  # Adjust max_workers as needed
        futures = [executor.submit(process_single_doc, doc) for doc in doc_files]
        for future in futures:
            embeddings, is_transcript = future.result()
            if embeddings:
                if is_transcript:
                    transcript_embeddings.extend(embeddings)
                else:
                    non_transcript_embeddings.extend(embeddings)

    print(f"\nTotal transcript embeddings: {len(transcript_embeddings)}")
    print(f"Total non-transcript embeddings: {len(non_transcript_embeddings)}")
    return transcript_embeddings, non_transcript_embeddings